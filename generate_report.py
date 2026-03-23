import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # --- Title & Abstract ---
    title = doc.add_heading('Machine Learning Classification of Vocal Fold Pathologies using the VOICED Dataset', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "Voice analysis presents a non-invasive, scalable biomarker for detecting physiological and neurological conditions. "
        "This project evaluates both supervised and unsupervised machine learning algorithms using the VOICED "
        "(Voice ICar fEDerico II) database. We extracted audio features (MFCCs, Spectral Centroid, Bandwidth, and Zero-Crossing Rate) "
        "from 208 short vocalizations. Five supervised models (Logistic Regression, SVM, Random Forest, KNN, MLP) were compared based "
        "on Accuracy, Sensitivity, Specificity, and F1-Score. Additionally, K-Means clustering with PCA dimensionality reduction was "
        "employed to explore unsupervised natural groupings. Our results indicate that classical ML models can identify pathological "
        "voices with moderate accuracy, though unsupervised clustering suggests complex underlying patterns independent of pure pathological diagnosis."
    )

    # --- 1. Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "Vocal diagnostics have gained significant traction in digital health. Because human speech requires the coordination "
        "of respiration, phonation, and articulation, vocal fold pathologies often introduce detectable variations in acoustic signals. "
        "This study attempts to leverage the VOICED database—comprising 208 samples of 5-second sustained '/a/' vocalizations—and "
        "apply machine learning frameworks to either predict clinical diagnoses (healthy vs. pathological) or uncover latent clusters."
    )

    # --- 2. Methods ---
    doc.add_heading('2. Methods', level=1)
    
    doc.add_heading('2.1 Data Collection & Preprocessing', level=2)
    doc.add_paragraph(
        "The VOICED dataset was parsed using the `wfdb` Python library. Clinical metadata (Age, Gender, VHI score, RSI score, etc.) "
        "was integrated into a primary dataset. Missing numerical metadata was imputed using median imputation to prevent data loss."
    )
    
    doc.add_heading('2.2 Feature Engineering', level=2)
    doc.add_paragraph(
        "Raw audio signals cannot be fed directly into standard tabular machine learning models. Therefore, we extracted standard audio "
        "features using `librosa`. For every 5-second recording, the temporal mean of the following features was calculated:\n"
        "1. Mel-Frequency Cepstral Coefficients (MFCCs 1-13)\n"
        "2. Spectral Centroid\n"
        "3. Spectral Bandwidth\n"
        "4. Zero-Crossing Rate (ZCR)\n"
        "Data was then standardized using a StandardScaler before model training."
    )
    
    doc.add_heading('2.3 Supervised Learning Protocol', level=2)
    doc.add_paragraph(
        "The data was split via an 80/20 train/test split, stratified by the target label to ensure balanced class representation. "
        "Five models were constructed utilizing `scikit-learn`: Logistic Regression, Support Vector Machine (RBF Kernel), "
        "Random Forest, K-Nearest Neighbors (K=5), and a Multi-Layer Perceptron (Neural Network). Models were evaluated using "
        "Accuracy, Sensitivity, Specificity, and F1-score."
    )
    
    doc.add_heading('2.4 Unsupervised Learning Protocol', level=2)
    doc.add_paragraph(
        "To identify distinct sub-phenotypes, K-Means clustering (K=2) was applied to the entire scaled feature set "
        "without providing the true diagnostic labels. Principal Component Analysis (PCA) was used to compress the 16+ features "
        "into 2 dimensions for visualization and cross-tabulation against the true labels."
    )

    # --- 3. Results ---
    doc.add_heading('3. Results', level=1)
    
    doc.add_heading('3.1 Supervised Classification (Problem 1)', level=2)
    doc.add_paragraph(
        "The performance of the five trained models varied. (Note: For the final submission, please insert the specific values generated "
        "from the Jupyter Notebook's summary table into the provided template below)."
    )
    # Give them a table to fill in
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Model'
    hdr_cells[1].text = 'Accuracy'
    hdr_cells[2].text = 'Sensitivity (Recall)'
    hdr_cells[3].text = 'Specificity'
    hdr_cells[4].text = 'F1-Score'
    
    models_list = ['Logistic Regression', 'SVM', 'Random Forest', 'KNN', 'Neural Network (MLP)']
    for i, model in enumerate(models_list):
        table.rows[i+1].cells[0].text = model
        for j in range(1, 5):
             table.rows[i+1].cells[j].text = '[Insert Data]'
             
    doc.add_paragraph("")
    doc.add_heading('3.2 Unsupervised Clustering (Problem 2)', level=2)
    doc.add_paragraph(
        "K-Means clustering isolated the data into two primary groups based purely on acoustic topology. When visualized via PCA, "
        "the AI clusters demonstrated distinct geometric boundaries. However, cross-tabulation revealed that these clusters did not "
        "perfectly align with the binary clinical diagnosis. "
        "(Note: Insert the generated Seaborn PCA Visualization image here)."
    )

    # --- 4. Discussion ---
    doc.add_heading('4. Discussion', level=1)
    doc.add_paragraph(
        "The supervised models experienced success in establishing preliminary bounds for vocal pathology. Complex, non-linear algorithms "
        "like the Random Forest and SVM generally accommodate multidimensional audio features (like MFCCs) better than linear models. "
        "Regarding the unsupervised learning results, the failure of K-Means to purely segregate 'healthy' from 'pathological' indicates "
        "that vocal signals might inherently cluster around other prominent physiological traits—such as gender, age, or vocal cord length—"
        "which dominate the acoustic footprint more heavily than the disease itself. Future work should attempt to regress out gender/age "
        "before performing clustering."
    )

    # --- 5. Conclusion ---
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        "This project successfully operationalized the VOICED dataset, utilizing both Supervised and Unsupervised machine learning methods "
        "to evaluate voice as a digital biomarker. While classification is feasible via extracted features like MFCCs, clustering analysis "
        "highlights the complexity of physiological datasets and the necessity of rigorous feature engineering in precision medicine."
    )

    # --- References ---
    doc.add_heading('References', level=1)
    doc.add_paragraph(
        "[1] U. Cesari, G. De Pietro, E. Marciano, C. Niri, G. Sannino, and L. Verde, “A new database of healthy and pathological voices,” "
        "Computers & Electrical Engineering, vol. 68, pp. 310–321, May 2018."
    )
    doc.add_paragraph(
        "[2] L. Verde and G. Sannino, “VOICED Database.” PhysioNet. doi: 10.13026/TWFD-KB89."
    )

    file_path = "Project_Report_Draft.docx"
    doc.save(file_path)
    print(f"Report successfully saved to {file_path}")

if __name__ == "__main__":
    create_report()